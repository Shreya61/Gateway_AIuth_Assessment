"""
Assembler Agent — builds publication-ready DOCX and PDF.

Structure:
  FRONT MATTER (roman numerals i–xii):
    Half-title | Title Page | Copyright | Dedication | Epigraph |
    TOC | Foreword | Preface | Acknowledgments

  BODY (Arabic numerals from Introduction):
    Introduction | Chapters

  BACK MATTER:
    Afterword | Appendix | Glossary | References (Further Reading) |
    About the Author | Back Cover Copy
"""

import os
import re
import shutil
import subprocess
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.logging.trace_logger import log_trace
from app.logging.cost_tracker import get_cost_summary
from app.utils.console import debug_print


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def page_break(doc):
    doc.add_page_break()


def soft_gap(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    return p


def centered(doc, text, size=28, bold=True, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    return p


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def set_page_numbering(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(existing)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)
    sect_pr.append(pg_num_type)


def add_page_number_footer(section, fmt="decimal"):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    add_field(p, " PAGE ")
    set_page_numbering(section, fmt=fmt, start=1)


def add_toc(doc):
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-2" \h \z \u')


def toc_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    return p


def add_visible_toc(doc, chapters, include_glossary=True, include_references=True):
    for item in ("Foreword", "Preface", "Acknowledgments", "Introduction"):
        toc_line(doc, item)
    for chapter in chapters:
        toc_line(doc, chapter_heading(chapter))
    for item in ("Afterword", "Appendix"):
        toc_line(doc, item)
    if include_glossary:
        toc_line(doc, "Glossary")
    if include_references:
        toc_line(doc, "Further Reading by Topic")
    toc_line(doc, "About the Author")
    toc_line(doc, "Back Cover")


def clean_chapter_title(chapter):
    title = str(chapter.get("title", "")).strip()
    title = re.sub(r"^chapter\s+\d+\s*:\s*", "", title, flags=re.IGNORECASE)
    title = re.sub(r"^chapter\s+\d+\s+", "", title, flags=re.IGNORECASE)
    return title or f"Chapter {chapter.get('chapter_number', '')}".strip()


def chapter_heading(chapter):
    return f"Chapter {chapter.get('chapter_number')}: {clean_chapter_title(chapter)}"


def safe_output_slug(state):
    requested = state.get("user_brief", {}).get("output_slug", "generated_book")
    slug = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(requested)).strip("._")
    return slug or "generated_book"


def subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)          # always 13pt — never varies
    run.font.name = "Times New Roman"
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def italic_center(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def write_section(doc, text_block):
    """Write a block of text, normalizing all heading variants to subheading style."""
    if not text_block:
        return
    for line in text_block.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Normalize ### and ## and # all to the same subheading style
        if line.startswith("### "):
            subheading(doc, line[4:])
        elif line.startswith("## "):
            subheading(doc, line[3:])
        elif line.startswith("# "):
            subheading(doc, line[2:])
        # Strip inline bold (**text**) markers so they don't appear as literal asterisks
        else:
            cleaned = line
            import re
            # Bold-only lines (e.g. **Section Title**) → treat as subheading
            bold_only = re.fullmatch(r"\*\*(.+)\*\*", cleaned)
            if bold_only:
                subheading(doc, bold_only.group(1))
            else:
                # Strip any remaining inline ** markers before writing body text
                cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
                cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
                body(doc, cleaned)


# ─────────────────────────────────────────────
# MAIN ASSEMBLER
# ─────────────────────────────────────────────

def run_assembler(state):
    try:
        os.makedirs("outputs", exist_ok=True)
        os.makedirs("logs", exist_ok=True)

        outline = state.get("outline", {})
        chapters = state.get("chapters", [])
        glossary = state.get("glossary", [])
        references = state.get("references", [])
        title = outline.get("book_title", "Untitled Book")
        subtitle = outline.get("subtitle", "")
        topic = state["user_brief"].get("topic", "")
        tone = state["user_brief"].get("tone", "")

        # Front/back matter
        dedication = state.get("dedication", "To every reader who made it past page one.")
        epigraph = state.get("epigraph", "")
        foreword = state.get("foreword", "")
        preface = state.get("preface", "")
        acknowledgments = state.get("acknowledgments", "")
        introduction = state.get("introduction", "")
        afterword = state.get("afterword", "")
        about_author = state.get("about_author", "")
        back_cover_copy = state.get("back_cover_copy", "")

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        add_page_number_footer(doc.sections[0], fmt="lowerRoman")

        # ── HALF TITLE ──────────────────────────────
        doc.add_paragraph("\n\n\n")
        centered(doc, title, size=22, bold=False)
        page_break(doc)

        # ── TITLE PAGE ───────────────────────────────
        doc.add_paragraph("\n\n")
        centered(doc, title, size=30)
        if subtitle:
            centered(doc, subtitle, size=16, bold=False, italic=True)
        doc.add_paragraph()
        centered(doc, "AIuthor Publishing", size=13, bold=False)
        page_break(doc)

        # ── COPYRIGHT ────────────────────────────────
        heading1(doc, "Copyright")
        body(doc, f"© AIuthor. All rights reserved.")
        body(doc, "ISBN: 978-0-00-000000-0 (placeholder)")
        body(doc, "First Edition, 2025")
        body(doc, "CIP Data: Available from the publisher.")
        body(doc, "No part of this publication may be reproduced without written permission.")
        soft_gap(doc)

        # ── DEDICATION ───────────────────────────────
        doc.add_paragraph("\n")
        italic_center(doc, dedication, size=13)
        soft_gap(doc)

        # ── EPIGRAPH ─────────────────────────────────
        if epigraph:
            doc.add_paragraph("\n")
            italic_center(doc, epigraph, size=12)
            soft_gap(doc)

        # ── TABLE OF CONTENTS ────────────────────────
        heading1(doc, "Table of Contents")
        add_toc(doc)
        add_visible_toc(doc, chapters, include_glossary=bool(glossary), include_references=bool(references))
        page_break(doc)

        # ── FOREWORD ─────────────────────────────────
        heading1(doc, "Foreword")
        write_section(doc, foreword or f"This book arrives at exactly the right moment. {topic} has never mattered more, and the author has found a way to make it accessible, urgent, and deeply human.")
        soft_gap(doc)

        # ── PREFACE ──────────────────────────────────
        heading1(doc, "Preface")
        write_section(doc, preface or f"Writing about {topic} is an act of optimism. This book exists because I believed readers deserved better than a manual.")
        soft_gap(doc)

        # ── ACKNOWLEDGMENTS ──────────────────────────
        heading1(doc, "Acknowledgments")
        write_section(doc, acknowledgments or "Thanks to everyone who read early drafts and told me what wasn't working.")

        body_section = doc.add_section(WD_SECTION.NEW_PAGE)
        add_page_number_footer(body_section, fmt="decimal")

        # ── INTRODUCTION (Arabic page numbers begin) ─
        heading1(doc, "Introduction")
        write_section(doc, introduction or f"This book is about {topic}. Not the abstract version, but the real version — the one that shows up in your decisions, your habits, and the quiet moments when you wonder if you're doing it right.")
        page_break(doc)

        # ── CHAPTERS ─────────────────────────────────
        for idx, chapter in enumerate(chapters):
            heading1(doc, chapter_heading(chapter))
            write_section(doc, chapter.get("content", ""))
            if idx < len(chapters) - 1:
                page_break(doc)

        # ── AFTERWORD ────────────────────────────────
        page_break(doc)
        heading1(doc, "Afterword")
        write_section(doc, afterword or "You made it. That alone means something. Now go do the thing.")

        # ── APPENDIX ─────────────────────────────────
        page_break(doc)
        heading1(doc, "Appendix")
        body(doc, "Key frameworks, checklists, and reference materials referenced throughout this book.")
        callback_index = state.get("callback_index", [])
        if callback_index:
            subheading(doc, "Callback & Motif Index")
            for cb in callback_index:
                if isinstance(cb, dict):
                    body(doc, f"• Chapter {cb.get('chapter','?')}: {cb.get('callback','')}")

        # ── GLOSSARY ─────────────────────────────────
        if glossary:
            page_break(doc)
            heading1(doc, "Glossary")
            for item in glossary:
                term = item.get("term", "")
                definition = item.get("definition", "")
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                bold_run = p.add_run(f"{term}: ")
                bold_run.bold = True
                bold_run.font.name = "Times New Roman"
                bold_run.font.size = Pt(12)
                def_run = p.add_run(definition)
                def_run.font.name = "Times New Roman"
                def_run.font.size = Pt(12)

        # ── REFERENCES ───────────────────────────────
        if references:
            page_break(doc)
            heading1(doc, "Further Reading by Topic")
            body(doc, "The following topical areas informed the research and ideas in this book:")
            for ref in references:
                body(doc, f"• {ref}")

        # ── ABOUT THE AUTHOR ─────────────────────────
        page_break(doc)
        heading1(doc, "About the Author")
        write_section(doc, about_author or "AIuthor is an agentic AI system designed to write books that humans actually want to read.")

        # ── BACK COVER COPY ───────────────────────────
        page_break(doc)
        heading1(doc, "Back Cover")
        write_section(doc, back_cover_copy or f"A book about {topic} for people who are done pretending they have it figured out.")

        # ── COST LEDGER ──────────────────────────────
        import json
        cost_summary = get_cost_summary()
        with open("logs/run_cost_summary.json", "w") as f:
            json.dump(cost_summary, f, indent=2)

        # ── SAVE DOCX ────────────────────────────────
        output_slug = safe_output_slug(state)
        docx_path = f"outputs/{output_slug}.docx"
        doc.save(docx_path)
        debug_print("DOCX SAVED:", docx_path)

        # ── GENERATE PDF via LibreOffice (if available) ──
        pdf_path = ""
        libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
        if libreoffice:
            result = subprocess.run(
                [libreoffice, "--headless", "--convert-to", "pdf",
                 "--outdir", "outputs/", docx_path],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode == 0:
                pdf_path = f"outputs/{output_slug}.pdf"
                debug_print("PDF SAVED:", pdf_path)
            else:
                debug_print("LibreOffice PDF conversion failed:", result.stderr)

        if not pdf_path:
            try:
                from reportlab.lib.pagesizes import LETTER
                from reportlab.pdfgen import canvas

                pdf_path = f"outputs/{output_slug}.pdf"
                c = canvas.Canvas(pdf_path, pagesize=LETTER)
                width, height = LETTER
                y = height - 72
                c.setFont("Times-Bold", 18)
                c.drawString(72, y, title[:80])
                y -= 36
                c.setFont("Times-Roman", 11)
                pdf_lines = [introduction] + [
                    f"Chapter {ch.get('chapter_number')}: {ch.get('title')}. {ch.get('content', '')}"
                    for ch in chapters
                ]
                for block in pdf_lines:
                    for line in str(block).replace("\n", " ").split(". "):
                        if y < 72:
                            c.showPage()
                            c.setFont("Times-Roman", 11)
                            y = height - 72
                        c.drawString(72, y, line[:95])
                        y -= 15
                c.save()
                debug_print("PDF SAVED:", pdf_path)
            except Exception as fallback_pdf_err:
                debug_print("Fallback PDF generation failed:", fallback_pdf_err)

        log_trace("assembler", "SUCCESS")

        return {
            **state,
            "generated_docx": docx_path,
            "generated_pdf": pdf_path,
        }

    except Exception as e:
        log_trace("assembler", "FAILED", str(e))
        debug_print("ASSEMBLER ERROR:", e)
        return {
            **state,
            "logs": state.get("logs", []) + [{"agent": "assembler", "error": str(e)}]
        }
